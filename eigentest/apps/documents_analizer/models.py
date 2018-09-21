# -*- coding: utf-8 -*-
""" apps documents_analizer models """

import os
from io import BytesIO

from django.db import models
from docx import Document as Docx

from utils.utils import highly_random_name, get_most_common_words


class Document(models.Model):
    """ Holds the document files """
    def update_filename(instance, filename):
        """ Changes the name of the file with a random name """
        path = os.path.join("documents_analizer", "documents")
        name = "{}{}".format(highly_random_name(),
                             os.path.splitext(filename)[1])
        return os.path.join(path, name)

    name = models.CharField(max_length=255)
    doc_file = models.FileField(upload_to=update_filename)

    class Meta:
        """ Meta class definitions """
        ordering = ['name']

    def __str__(self):
        """ Returns a string representation """
        return self.name

    def get_doc_file_extension(self):
        """ Returns the lowercase extension of the doc_file """
        return os.path.splitext(self.doc_file.name)[-1][1:].lower()

    def get_full_text_from_source(self):
        """ Returns the full text from the source """
        extension = self.get_doc_file_extension()

        if extension in ('txt', ''):
            # string = unicode(string)
            return self.doc_file.read().decode("utf-8")
        elif extension == 'docx':
            docx_document = Docx(BytesIO(self.doc_file.read()))
            return "\n".join(p.text for p in docx_document.paragraphs)
        elif extension == 'pdf':
            raise NotImplementedError()
        else:
            raise ValueError("file_format not supported")

    def get_lines_from_source(self):
        """ Returns a list with the lines from the source """
        extension = self.get_doc_file_extension()
        if extension in ('txt', ''):
            return tuple(line.decode('utf-8') for line in self.doc_file.readlines())
        elif extension == 'docx':
            docx_document = Docx(BytesIO(self.doc_file.read()))
            return tuple(paragrah.text for paragrah in docx_document.paragraphs)
        elif extension == 'pdf':
            raise NotImplementedError()
        else:
            raise ValueError("file_format not supported")

    def extract_hashtags(self):
        """
        Extracts the most common hashtags from the document and tries to add
        them to the Hashtag table.
        """
        most_common_words = get_most_common_words(
            self.get_full_text_from_source())
        queryset = Hashtag.objects.all()
        for word in most_common_words:
            tag = queryset.get_or_create(word=word.lower())[0]
            tag.documents.add(self)


class Hashtag(models.Model):
    """ Holds the hashtags and the related documents """
    word = models.CharField(max_length=60, unique=True)
    documents = models.ManyToManyField(Document)

    class Meta:
        """ Meta class definitions """
        ordering = ['word']

    def __str__(self):
        """ Returns a string representation """
        return self.word
